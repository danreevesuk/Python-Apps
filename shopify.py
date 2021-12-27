import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import requests
import json 
from docx2pdf import convert
from docx import Document



    
def main(item, name):    

    doc = docx.Document() 
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)



    Main_Heading = doc.add_heading().add_run("                                                        Master Title Deed")
    Main_Heading.alignment = 1
    font = Main_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(14)    



    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    starting1 = doc.add_paragraph("Deed of Change of Name and Title (Deed Poll)")
    starting1.alignment = 1
    starting1.paragraph_format.space_before = Pt(0)
    starting1.paragraph_format.space_after = Pt(0)



    starting2 = doc.add_paragraph("of")
    starting2.alignment = 1
    starting2.paragraph_format.space_before = Pt(0)
    starting2.paragraph_format.space_after = Pt(0)


    starting3 = doc.add_paragraph(name)
    starting3.alignment = 1
    starting3.paragraph_format.space_before = Pt(0)
    starting3.paragraph_format.space_after = Pt(0)


    starting4 = doc.add_paragraph("now")
    starting4.alignment = 1
    starting4.paragraph_format.space_before = Pt(0)
    starting4.paragraph_format.space_after = Pt(0)


    starting5 = doc.add_paragraph(item +" " + name)
    starting5.alignment = 1
    starting5.paragraph_format.space_before = Pt(0)
    starting5.paragraph_format.space_after = Pt(1)


    starting6 = doc.add_paragraph("BY THIS DEED OF CHANGE OF NAME AND TITLE made by myself the undersigned "+ item+" "+ name +" of")
    starting6.alignment = 0
    starting6.space_before = Pt(0)



    Head_Break=doc.add_paragraph("____________________________________________________________________________________")
    Head_Break.space_before = Pt(0)
    Head_Break.alignment = 1

    starting7 = doc.add_paragraph("Address")
    starting7.alignment = 1
    starting7.space_before = Pt(0)

    starting8 = doc.add_paragraph("HEREBY DECLARE AS FOLLOWS:")
    starting8.alignment = 0
    starting8.space_before = Pt(0)


    starting9 = doc.add_paragraph("1. 	I ABSOLUTELY and entirely renounce, relinquish and abandon the use of my former name of  "+name+" and assume, adopt and determine to take and use from the date hereof the name and title of " + item + " "+ name+" in substitution for my former name of "+name+".")
    starting9.alignment = 0
    starting9.space_before = Pt(0)


    starting10 = doc.add_paragraph("2. 	I SHALL at all times hereafter in all records, deeds, documents and other writings and in all transactions and on all occasions whatsoever use and subscribe the said name and title of " + item + " " +name+" as my name in substitution for my former name of "+name+" so relinquished as aforesaid to the intent that I may hereafter be called, known or distinguished by the name and title of " +item+ " " +name+" only and not by my former name of "+name+".")
    starting10.alignment = 0
    starting10.space_before = Pt(0)


    starting11 = doc.add_paragraph("3. 	I AUTHORISE and require all persons at all times to describe, designate and address me by my adopted name and title of " +item+ " "+name+".")
    starting11.alignment = 0
    starting11.space_before = Pt(0)


    starting12 = doc.add_paragraph("IN WITNESS whereof I have hereunto subscribed my adopted and substituted name and title of " +item+ " " +name+" and also my former name of "+name+".")
    starting12.alignment = 0
    starting12.space_before = Pt(0)


    starting13 = doc.add_paragraph("Dated this ________ day of ________ in the year ______ ")
    starting13.alignment = 0
    starting13.space_before = Pt(0)
    starting13.space_after = Pt(0)



    table = doc.add_table(rows=0, cols=2)
    table.alignment = 1 

    row0=table.add_row().cells
    p=row0[0].add_paragraph('SIGNED AS A DEED AND DELIVERED \n by the above named')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p=row0[1].add_paragraph('In the presence of:')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


    p.alignment=WD_ALIGN_PARAGRAPH.LEFT


    row1=table.add_row().cells
    p=row1[0].add_paragraph('_____________________________')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=row1[1].add_paragraph('Name ___________________________________')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    row2=table.add_row().cells
    p=row2[0].add_paragraph(item+ " "+name)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=row2[1].add_paragraph("Witness's signature ________________________")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    row3=table.add_row().cells
    p=row3[0].add_paragraph("Formerly known as")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p=row3[1].add_paragraph("Address _________________________________")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


    row4=table.add_row().cells
    p=row4[0].add_paragraph('_____________________________')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p=row4[1].add_paragraph('_________________________________________')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


    row5=table.add_row().cells
    p=row5[0].add_paragraph(name)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p=row5[1].add_paragraph('Occupation ______________________________')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


    doc.save(order_number+" "+name+'.docx')
    inputFile  = order_number+" "+name+".docx"
    convert(inputFile)

def generate_certificate(file):
    doc = Document(file)
    for p in doc.paragraphs:
        if 'Name' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'Name' in inline[i].text:
                    text = inline[i].text.replace('Name', name)
                    inline[i].text = text
            print (p.text)

    doc.save("Certificate "+order_number+" "+name+'.docx')
    inputFile_cert  = "Certificate " +order_number+" "+name+".docx"
    convert(inputFile_cert)


    return 1

while True:
    time.sleep(300)

    res = requests.get('https://8fa5930845f0b60d4b07950dc66ba546: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@prestige-titles.myshopify.com/admin/api/2021-10/orders.json?limit=1')
    # res = requests.get('https://8fa5930845f0b60d4b07950dc66ba546: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@prestige-titles.myshopify.com/admin/api/2021-10/orders/3805879926993.json')

    response = json.loads(res.text)

    product = response['orders'][0]['line_items'][0]['name']
    # product = response['order']['line_items'][0]['name']
    print(product)
    order_number = response['orders'][0]['order_number']
    order_number = str(order_number)


    if 'Lord Title' in product:
        item = 'Lord'
        name = response['orders'][0]['line_items'][0]['properties'][0]['value']
        name = name.title()
        main(item, name)
        generate_certificate("lord.docx")
        print("Lord")
    elif 'Lady Title' in product:
        item = 'Lady'
        name = response['orders'][0]['line_items'][0]['properties'][0]['value']
        name = name.title()
        generate_certificate("lady.docx")
        main(item, name)
        print("Lady")
    elif 'Lord & Lady' in product:
        item = 'Lady'
        name = response['orders'][0]['line_items'][0]['properties'][0]['value']
        name = name.title()
        print(name)
        main(item , name)
        generate_certificate("lady.docx")

        
        item = 'Lord'
        name = response['orders'][0]['line_items'][0]['properties'][1]['value']
        name = name.title()
        print(name)
        main(item, name)
        generate_certificate("lord.docx")



