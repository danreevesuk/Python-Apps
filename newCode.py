import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import requests
import json 
from docx2pdf import convert
from docx import Document
from PyPDF2 import PdfFileWriter, PdfFileReader
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import textwrap


res = requests.get('https://8fa5930845f0b60d4b07950dc66ba546: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@prestige-titles.myshopify.com/admin/api/2021-10/orders.json?fulfillment_status=unfullfilled&limit=200')


response = json.loads(res.text)
product = response['orders']

    
def main(item, name):    

    doc = docx.Document() 
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_height  = Inches(11.7)
        section.page_width = Inches(8.3)



    Main_Heading = doc.add_heading().add_run("                                                       Master Title Deed")
    Main_Heading.alignment = 1
    font = Main_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(14)    



    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    starting1 = doc.add_paragraph("Deed of Change of Name and Title (Deed Poll)")
    starting1.alignment = 1
    starting1.paragraph_format.space_before = Pt(0)
    starting1.paragraph_format.space_after = Pt(0)



    starting2 = doc.add_paragraph("of")
    starting2.alignment = 1
    starting2.paragraph_format.space_before = Pt(0)
    starting2.paragraph_format.space_after = Pt(0)


    starting3 = doc.add_paragraph(name)
    starting3.alignment = 1
    starting3.paragraph_format.space_before = Pt(0)
    starting3.paragraph_format.space_after = Pt(0)


    starting4 = doc.add_paragraph("now")
    starting4.alignment = 1
    starting4.paragraph_format.space_before = Pt(0)
    starting4.paragraph_format.space_after = Pt(0)


    starting5 = doc.add_paragraph(item +" " + name)
    starting5.alignment = 1
    starting5.paragraph_format.space_before = Pt(0)
    starting5.paragraph_format.space_after = Pt(1)


    starting6 = doc.add_paragraph("BY THIS DEED OF CHANGE OF NAME AND TITLE made by myself the undersigned "+ item+" "+ name +" of")
    starting6.alignment = 0
    starting6.space_before = Pt(0)



    Head_Break=doc.add_paragraph("_________________________________________________________________________________")
    Head_Break.space_before = Pt(0)
    Head_Break.space_after = Pt(0)
    Head_Break.alignment = 1

    starting7 = doc.add_paragraph("Address")
    starting7.alignment = 1
    starting7.space_before = Pt(0)
    starting7.space_after = Pt(0)


    starting8 = doc.add_paragraph("HEREBY DECLARE AS FOLLOWS:")
    starting8.alignment = 0
    starting8.space_before = Pt(0)
    starting8.space_after = Pt(0)



    starting9 = doc.add_paragraph("1. 	I ABSOLUTELY and entirely renounce, relinquish and abandon the use of my former name of  "+name+" and assume, adopt and determine to take and use from the date hereof the name and title of " + item + " "+ name+" in substitution for my former name of "+name+".")
    starting9.alignment = 0
    starting9.space_before = Pt(0)


    starting10 = doc.add_paragraph("2. 	I SHALL at all times hereafter in all records, deeds, documents and other writings and in all transactions and on all occasions whatsoever use and subscribe the said name and title of " + item + " " +name+" as my name in substitution for my former name of "+name+" so relinquished as aforesaid to the intent that I may hereafter be called, known or distinguished by the name and title of " +item+ " " +name+" only and not by my former name of "+name+".")
    starting10.alignment = 0
    starting10.space_before = Pt(0)


    starting11 = doc.add_paragraph("3. 	I AUTHORISE and require all persons at all times to describe, designate and address me by my adopted name and title of " +item+ " "+name+".")
    starting11.alignment = 0
    starting11.space_before = Pt(0)


    starting12 = doc.add_paragraph("IN WITNESS whereof I have hereunto subscribed my adopted and substituted name and title of " +item+ " " +name+" and also my former name of "+name+".")
    starting12.alignment = 0
    starting12.space_before = Pt(0)


    starting13 = doc.add_paragraph("Dated this ________ day of ________ in the year ______ ")
    starting13.alignment = 0
    starting13.space_before = Pt(0)
    starting13.space_after = Pt(0)



    table = doc.add_table(rows=0, cols=2)
    table.alignment = 1 

    row0=table.add_row().cells
    p=row0[0].add_paragraph('SIGNED AS A DEED AND DELIVERED \n by the above named')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p=row0[1].add_paragraph('In the presence of:')
    p.paragraph_format.space_before = Pt(0) 
    p.paragraph_format.space_after = Pt(0)


    p.alignment=WD_ALIGN_PARAGRAPH.LEFT


    row1=table.add_row().cells
    p=row1[0].add_paragraph('_____________________________\n'+item+ " "+name)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=row1[1].add_paragraph('Name ___________________________________')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    row2=table.add_row().cells
    p=row2[0].add_paragraph("Formerly known as")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p=row2[1].add_paragraph("Witness's signature ________________________")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    row3=table.add_row().cells
    p=row3[0].add_paragraph('_____________________________\n'+name)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=row3[1].add_paragraph("Address _________________________________\n\n_________________________________________")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


    # row4=table.add_row().cells
    # p=row4[0].add_paragraph(" ")
    # p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # p.paragraph_format.space_before = Pt(0)
    # p.paragraph_format.space_after = Pt(0)
    # p=row4[1].add_paragraph(' ')
    # p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # p.paragraph_format.space_before = Pt(0)
    # p.paragraph_format.space_after = Pt(0)


    row5=table.add_row().cells
    p=row5[0].add_paragraph(" ")
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p=row5[1].add_paragraph('Occupation ______________________________')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


    doc.save("generated_files/"+order_number+" "+name+'.docx')
    inputFile  = "generated_files/"+order_number+" "+name+".docx"
    convert(inputFile)



def generate_certificate():
    buffer = BytesIO()


    pdfmetrics.registerFont(TTFont('times new roman', 'times new roman.ttf'))
    pdfmetrics.registerFont(TTFont('times new roman italic', 'times new roman italic.ttf'))

    p = canvas.Canvas(buffer, pagesize=A4)

    p.setFont('times new roman italic', 32)
    x_pos = 300
    y_pos = 410
    y_offset = 40
    if len(item+name) > 25:
            wraps = textwrap.wrap(item+" "+name, 25)
            for x in range(len(wraps)):
                p.drawCentredString(x_pos, y_pos, wraps[x])
                y_pos -= y_offset
            y_pos += y_offset
    else:
        p.drawCentredString(x=300,y=390, text=item+" "+name)



    p.showPage()
    p.save()

    #move to the beginning of the StringIO buffer
    buffer.seek(0)
    newPdf = PdfFileReader(buffer)


    existingPdf = PdfFileReader(open('lord.pdf', 'rb'))
    output = PdfFileWriter()
    # add the "watermark" (which is the new pdf) on the existing page
    page = existingPdf.getPage(0)
    page.mergePage(newPdf.getPage(0))
    output.addPage(page)
    # finally, write "output" to a real file
    outputStream = open("generated_files/"+"Certificate " +order_number+" "+name+".pdf", 'wb')
    output.write(outputStream)
    outputStream.close()


    return 1


for i in range(len(product)):

    find_product = response['orders']
    new  = find_product[i]
    product = new['line_items'][0]['name']
    

    # product = response['order']['line_items'][0]['name']
    print(product)
    order_number = new['order_number']
    order_number = str(order_number)


    if 'Lord Title' in product:
        item = 'Lord'
        name = new['line_items'][0]['properties'][0]['value']
        name = name.title().rstrip()
        main(item, name)
        generate_certificate()
        print(item)
        print(name)
    elif 'Lady Title' in product:
        item = 'Lady'
        name = new['line_items'][0]['properties'][0]['value']
        name = name.title().rstrip()
        main(item, name)
        generate_certificate()
        print(item)
        print(name)
    elif 'Lord & Lady' in product:
        item = 'Lady'
        name = new['line_items'][0]['properties'][0]['value']
        name = name.title().rstrip()
        main(item, name)
        generate_certificate()
        print(item)
        print(name)

        
        item = 'Lord'
        name = new['line_items'][0]['properties'][1]['value']
        name = name.title().rstrip()
        main(item, name)
        generate_certificate()
        print(item)
        print(name)

    elif 'Baron Title' in product:
        item = 'Baron'
        name = new['line_items'][0]['properties'][0]['value']
        name = name.title().rstrip()
        main(item, name)
        generate_certificate()
        print(item)
        print(name)
    elif 'Baroness Title' in product:
        item = 'Baroness'
        name = new['line_items'][0]['properties'][0]['value']
        name = name.title().rstrip()
        main(item, name)
        generate_certificate()
        print(item)
        print(name)
    elif 'Baron & Baroness' in product:
        item = 'Baroness'
        name = new['line_items'][0]['properties'][0]['value']
        name = name.title().rstrip()
        main(item, name)
        generate_certificate()
        print(item)
        print(name)

        
        item = 'Baron'
        name = new['line_items'][0]['properties'][1]['value']
        name = name.title().rstrip()
        main(item, name)
        generate_certificate()
        print(item)
        print(name)


    else:
        try:
            product2 = new['line_items'][1]['name']


            if 'Lord Title' in product2:
                item = 'Lord'
                name = new['line_items'][1]['properties'][0]['value']
                name = name.title().rstrip()
                main(item, name)
                generate_certificate()
                print(item)
                print(name)
            elif 'Lady Title' in product2:
                item = 'Lady'
                name = new['line_items'][1]['properties'][0]['value']
                name = name.title().rstrip()
                generate_certificate()
                print(item)
                print(name)
            elif 'Lord & Lady' in product2:
                item = 'Lady'
                name = new['line_items'][1]['properties'][0]['value']
                name = name.title().rstrip()
                generate_certificate()
                print(item)
                print(name)

                
                item = 'Lord'
                name = new['line_items'][1]['properties'][1]['value']
                name = name.title().rstrip()
                generate_certificate()
                print(item)
                print(name)


            elif 'Baron Title' in product2:
                item = 'Baron'
                name = new['line_items'][1]['properties'][0]['value']
                name = name.title().rstrip()
                generate_certificate()
                print(item)
                print(name)
            elif 'Baroness Title' in product2:
                item = 'Baroness'
                name = new['line_items'][1]['properties'][0]['value']
                name = name.title().rstrip()
                generate_certificate()
                print(item)
                print(name)
                print("Baroness")
            elif 'Baron & Baroness' in product2:
                item = 'Baroness'
                name = new['line_items'][1]['properties'][0]['value']
                name = name.title().rstrip()
                generate_certificate()
                print(item)
                print(name)

                
                item = 'Baron'
                name = new['line_items'][1]['properties'][1]['value']
                name = name.title().rstrip()
                generate_certificate()
                print(item)
                print(name)

        except:
            continue
        

    i = i+1
    time.sleep(5)

