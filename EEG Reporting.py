import tempfile
import os
from tkinter import *
import tkinter.font as font
import tkinter.messagebox
from datetime import date
import docx
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_LINE_SPACING


#=======================================================================


doc = docx.Document() 
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

#===========================================================================
root = Tk()
root.geometry("820x750+300+0")
root.title("EEG Software")
root.configure(background= "Cadet Blue")

MenuFrame = Frame(root, bg= "Cadet Blue", bd=5 , relief=RIDGE)
MenuFrame.pack()

FrameTitle = Frame (MenuFrame, bg= "powder Blue", bd=5 , relief=RIDGE)
FrameTitle.pack(side=TOP)

FrameInfo = Frame(MenuFrame, bg= "powder Blue", bd=5 , relief=RIDGE)
FrameInfo.pack(side=TOP)

FrameHistory = Frame(MenuFrame, bg= "powder Blue", bd=5 , relief=RIDGE)
FrameHistory.pack(side=TOP)

Reciept = Frame(MenuFrame, bg= "powder Blue", bd=5 , relief=RIDGE)
Reciept.pack(side=TOP)

Buttons = Frame(MenuFrame, bg= "powder Blue", bd=5 , relief=RIDGE)
Buttons.pack(side=BOTTOM)
#================================================================================


def dPrint(): 
    doc.add_picture('eeg.png')
    Main_Heading = doc.add_heading().add_run("                              ELECTROENCEPHALOGRAM")
    Main_Heading.alignment = 1
    font = Main_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(18)    
    
    
    Head_Break=doc.add_paragraph("____________________________________________________________________________________")
    Head_Break.space_before = Pt(0)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    table = doc.add_table(rows=3, cols=2)
    table.alignment = 1 
    
    a1 = Patient_Name.cget("text")
    a2 = Patient_Name_txt.get()
    a2 = a2.upper()
    d1 = Date.cget("text")
    d2 = Date_txt.get()
    cell1 = table.cell(0, 0)
    cell1.text = a1+ ": "+a2
    cell2 = table.cell(0, 1)
    cell2.text = "            "+d1+ ": "+d2
    
    
    
    b1 = Patient_Age.cget("text")
    b2 = Patient_Age_txt.get()
    b2 = b2.upper()
    e1 = EEG.cget("text")
    e2 = EEG_txt.get()
    cell3 = table.cell(1, 0)
    cell3.text = b1+ ": "+b2
    cell4 = table.cell(1, 1)
    cell4.text = "            "+e1+ ": "+e2
    
    
    c1 = Patient_Sex.cget("text")
    c2 = Patient_Sex_txt.get()
    c2 = c2.upper()
    f1 = Ref_Doctor.cget("text")
    f2 = Ref_Doctor_txt.get()
    f2 = f2.upper()
    cell5 = table.cell(2, 0)
    cell5.text = c1+ ": "+c2
    cell6 = table.cell(2, 1)
    cell6.text = "            "+f1+ ": "+"DR. "+f2
    
    
    Cell_Break=doc.add_paragraph("____________________________________________________________________________________")
    
    Hist_Heading = doc.add_heading().add_run("PATIENT'S HISTORY:")
    Hist_Heading.underline = True
    font = Hist_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(13)   
    
    
    Hist_Para_Text = History_txt.get(1.0,"end-1c")
    Hist_Para = doc.add_paragraph(Hist_Para_Text)
    #Hist_Break=doc.add_paragraph("____________________________________________________________________________________")
    
    
    Doc_Heading = doc.add_heading().add_run("DOCTOR'S REPORT:")
    Doc_Heading.underline = True
    font = Doc_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(13) 
    
    Doc_Para_Text = Doc_report_txt.get(1.0,"end-1c")
    Doc_Para = doc.add_paragraph(Doc_Para_Text)
    Doc_Para.alignment = 3
    #Doc_Break=doc.add_paragraph("____________________________________________________________________________________")    
    
    
    Imp_Heading = doc.add_heading().add_run("IMPRESSION:")
    Imp_Heading.underline = True
    font = Imp_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(13) 
    
    Imp_Para_Text = Impression_txt.get(1.0,"end-1c")
    Imp_Para = doc.add_paragraph(Imp_Para_Text)
    #Imp_Break=doc.add_paragraph("____________________________________________________________________________________")
    
    
    Conc_Heading = doc.add_heading().add_run("CONCLUSION:")
    Conc_Heading.underline = True
    font = Conc_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(13) 
    
    
    Conc_Para_Text = Conslusion_txt.get(1.0,"end-1c")
    Conc_Para = doc.add_paragraph(Conc_Para_Text)
    Conc_Break=doc.add_paragraph("____________________________________________________________________________________")
    Rep_Para_Text = reported_by_txt.get(1.0,"end-1c")
    Rep_Para = doc.add_paragraph(Rep_Para_Text)
    #Rep_Break=doc.add_paragraph("____________________________________________________________________________________")
    

    
    doc.save('EEG PRINTING.docx') 
    os.startfile("EEG PRINTING.docx", "print")#<---'print'
    os.remove("EEG PRINTING.docx")
  
    
#========================================================================================    
    
def Open(): 
    doc.add_picture('eeg.png')
    Main_Heading = doc.add_heading().add_run("                              ELECTROENCEPHALOGRAM")
    Main_Heading.alignment = 1
    font = Main_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(18)    
    
    
    Head_Break=doc.add_paragraph("____________________________________________________________________________________")
    Head_Break.space_before = Pt(0)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    table = doc.add_table(rows=3, cols=2)
    table.alignment = 1 
    
    a1 = Patient_Name.cget("text")
    a2 = Patient_Name_txt.get()
    a2 = a2.upper()
    d1 = Date.cget("text")
    d2 = Date_txt.get()
    cell1 = table.cell(0, 0)
    cell1.text = a1+ ": "+a2
    cell2 = table.cell(0, 1)
    cell2.text = "            "+d1+ ": "+d2
    
    
    
    b1 = Patient_Age.cget("text")
    b2 = Patient_Age_txt.get()
    b2 = b2.upper()
    e1 = EEG.cget("text")
    e2 = EEG_txt.get()
    cell3 = table.cell(1, 0)
    cell3.text = b1+ ": "+b2
    cell4 = table.cell(1, 1)
    cell4.text = "            "+e1+ ": "+e2
    
    
    c1 = Patient_Sex.cget("text")
    c2 = Patient_Sex_txt.get()
    c2 = c2.upper()
    f1 = Ref_Doctor.cget("text")
    f2 = Ref_Doctor_txt.get()
    f2 = f2.upper()
    cell5 = table.cell(2, 0)
    cell5.text = c1+ ": "+c2
    cell6 = table.cell(2, 1)
    cell6.text = "            "+f1+ ": "+"DR. "+f2
    
    
    Cell_Break=doc.add_paragraph("____________________________________________________________________________________")
    
    Hist_Heading = doc.add_heading().add_run("PATIENT'S HISTORY:")
    Hist_Heading.underline = True
    font = Hist_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(13)   
    
    
    Hist_Para_Text = History_txt.get(1.0,"end-1c")
    Hist_Para = doc.add_paragraph(Hist_Para_Text)
    #Hist_Break=doc.add_paragraph("____________________________________________________________________________________")
    
    
    Doc_Heading = doc.add_heading().add_run("DOCTOR'S REPORT:")
    Doc_Heading.underline = True
    font = Doc_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(13) 
    
    Doc_Para_Text = Doc_report_txt.get(1.0,"end-1c")
    Doc_Para = doc.add_paragraph(Doc_Para_Text)
    Doc_Para.alignment = 3
    #Doc_Break=doc.add_paragraph("____________________________________________________________________________________")    
    
    
    Imp_Heading = doc.add_heading().add_run("IMPRESSION:")
    Imp_Heading.underline = True
    font = Imp_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(13) 
    
    Imp_Para_Text = Impression_txt.get(1.0,"end-1c")
    Imp_Para = doc.add_paragraph(Imp_Para_Text)
    #Imp_Break=doc.add_paragraph("____________________________________________________________________________________")
    
    
    Conc_Heading = doc.add_heading().add_run("CONCLUSION:")
    Conc_Heading.underline = True
    font = Conc_Heading.font
    font.color.rgb = RGBColor(0,0,0)
    font.size= Pt(13) 
    
    
    Conc_Para_Text = Conslusion_txt.get(1.0,"end-1c")
    Conc_Para = doc.add_paragraph(Conc_Para_Text)
    Conc_Break=doc.add_paragraph("____________________________________________________________________________________")
    Rep_Para_Text = reported_by_txt.get(1.0,"end-1c")
    Rep_Para = doc.add_paragraph(Rep_Para_Text)
    #Rep_Break=doc.add_paragraph("____________________________________________________________________________________")
    

    
    doc.save('EEG PRINTING.docx') 
    os.startfile("EEG PRINTING.docx")#<---'print'
    
    
    
#===================================================================================    
    
def Reset():
    Reset = tkinter.messagebox.askyesno("Attention", "Do You Want to Reset All ?")
    if Reset==1:
        Patient_Name_txt.delete(0, END)
        Patient_Age_txt.delete(0, END)
        Patient_Sex_txt.delete(0, END)
        #Date_txt.delete(0, END)
        EEG_txt.delete(0, END)
        Ref_Doctor_txt.delete(0, END)
        History_txt.delete(1.0, END)
        Doc_report_txt.delete(1.0, END)
        Impression_txt.delete(1.0, END)
        Conslusion_txt.delete(1.0, END)
        reported_by_txt.delete(1.0, END)
        os.remove("EEG PRINTING.docx")
#========================================================================================
        
def iExit():
    iExit = tkinter.messagebox.askyesno("Attention", "Do You Want to Quit ?")
    if iExit==1:
        root.destroy()
        return


    
    
        

#================================================================================
lbTitle = Label(FrameTitle, width= 40, bg= "white", fg="black", bd=4 ,font=('arial',24,'bold'), text="EEG REPORTING")
lbTitle.grid(row =0 , column=0)

#================================================================================

Patient_Name= Label(FrameInfo, width= 13,height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',11), text="PATIENT'S NAME")
Patient_Name.grid(row=0, column=0, sticky='n', ipadx= 0)

Patient_Name_txt = Entry(FrameInfo, width= 25, bg= "white" , bd=5,font=('arial',12))
Patient_Name_txt.grid(row =0 , column=1)

Patient_Age= Label(FrameInfo, width= 13, height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',12), text="AGE")
Patient_Age.grid(row=1, column=0, sticky='n')

Patient_Age_txt = Entry(FrameInfo, width= 25, bg= "white" , bd=5,font=('arial',12))
Patient_Age_txt.grid(row =1 , column=1)

Patient_Sex= Label(FrameInfo, width= 13, height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',12), text="SEX")
Patient_Sex.grid(row=2, column=0, sticky='n')

Patient_Sex_txt= Entry(FrameInfo, width= 25, bg= "white" , bd=5,font=('arial',12))
Patient_Sex_txt.grid(row=2, column=1,)

Date= Label(FrameInfo, width= 13,height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',12), text="DATE")
Date.grid(row=0, column=2, sticky='n', ipadx= 0)

Date_txt = Entry(FrameInfo, width= 25, bg= "white" , bd=5,font=('arial',12))
today = date.today()
Date_txt.insert(10,today.strftime("%d/%m/%Y"))
Date_txt.grid(row =0 , column=3)

EEG= Label(FrameInfo, width= 13, height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',12), text="EEG #")
EEG.grid(row=1, column=2, sticky='n')

EEG_txt = Entry(FrameInfo, width= 25, bg= "white" , bd=5,font=('arial',12))
EEG_txt.grid(row =1 , column=3)

Ref_Doctor= Label(FrameInfo, width= 13,height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',12), text="REF. BY")
Ref_Doctor.grid(row=2, column=2, sticky='n')

Ref_Doctor_txt = Entry(FrameInfo, width= 25, bg= "white" , bd=5,font=('arial',12))
Ref_Doctor_txt.grid(row =2 , column=3)


#=====================================================================================================

Patient_History = Label(FrameHistory, width= 42,height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',22, "bold"), text="Patient's History")
Patient_History.grid(row=0, column=0, sticky='w')

History_txt = Text(FrameHistory, width= 84, height = 1, bg= "white", bd=4 ,font=('arial',12))
History_txt.grid(row =1 , column=0)

#=====================================================================================================

Doc_report = Label(FrameHistory, width= 42,height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',22, "bold"), text="DOCTOR'S REPORT")
Doc_report.grid(row=2, column=0, sticky='w')

Doc_report_txt = Text(FrameHistory, width= 84, height = 7, bg= "white", bd=4 ,font=('arial',12,))
Doc_report_txt.grid(row =3 , column=0)

#=====================================================================================================

Patient_Impression = Label(FrameHistory, width= 42,height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',22, "bold"), text="Impression")
Patient_Impression.grid(row=4, column=0, sticky='w')

Impression_txt = Text(FrameHistory, width= 84, height = 1, bg= "white", bd=4 ,font=('arial',12,))
Impression_txt.grid(row =5 , column=0)

#=====================================================================================================

Patient_Conclusion = Label(FrameHistory, width= 42,height = 1, bg= "powder Blue", fg="black", bd=5 ,font=('arial',22, "bold"), text="Conclusion")
Patient_Conclusion.grid(row=6, column=0, sticky='w')

Conslusion_txt = Text(FrameHistory, width= 84, height = 3, bg= "white", bd=4 ,font=('arial',12,))
Conslusion_txt.grid(row =7 , column=0)

#=====================================================================================================

reported_by = Label(FrameHistory, width= 10,height = 3, bg= "powder Blue", fg="black", bd=5 ,font=('arial',12, "bold"), text="Consultant")
reported_by.grid(row=8, column=0,sticky='w')

reported_by_txt = Text(FrameHistory, width= 65, height = 4, bg= "white", bd=5 ,font=('arial',10))
reported_by_txt.grid(row =8 , column=0, sticky='e')
reported_by_txt.insert(1.0,"  DR. IRFAN HASHMAT\n     (M.B.B.S, F.P.C.S)\nConsultant Neurologist")

#=====================================================================================================


btn1 = Button(Buttons, font=("arial",16,"bold"), width = 14, text="Print", command = dPrint)
btn1.grid(row =1 , column=0)

btn4 = Button(Buttons, font=("arial",16,"bold"), width = 14, text="Open", command = Open)
btn4.grid(row =1 , column=1)

btn2 = Button(Buttons, font=("arial",16,"bold"), width = 14, text="Reset All" , command = Reset)
btn2.grid(row =1 , column=2)

btn3 = Button(Buttons, font=("arial",16,"bold"), width = 14, text="Exit", command = iExit)
btn3.grid(row =1 , column=3)

root.mainloop()
